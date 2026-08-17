from __future__ import annotations

import re
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables"
ASSETS = ROOT / "generated_assets_v2"
OUT.mkdir(parents=True, exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)

# Base: narrative_proposal preset. Named academic-journal overrides are reused
# throughout: Times New Roman / SimSun body, restrained ink-blue headings,
# compact title block, and journal-review line spacing.
INK = "20354A"
BLUE = "2E74B5"
TEAL = "177E89"
GOLD = "B7791F"
MUTED = "677788"
PALE_BLUE = "EAF2F8"
PALE_TEAL = "E8F4F2"
PALE_GOLD = "FFF4DA"
PALE_GRAY = "F2F5F7"
PALE_RED = "FCEBE8"
WHITE = "FFFFFF"
BLACK = "111111"


def _set_run_font(run, name: str, size: float | None = None, bold=None,
                  italic=None, color: str | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _cell_margins(cell, top=90, start=105, bottom=90, end=105) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def _add_hyperlink(paragraph, text: str, url: str, font_name: str, size: float) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), font_name)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.extend([rfonts, color, underline, sz])
    new_run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(https?://\S+|\*\*.*?\*\*|\*.*?\*|`.*?`)")


def _inline(paragraph, text: str, font_name: str, size: float) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            _set_run_font(paragraph.add_run(text[cursor:match.start()]), font_name, size)
        token = match.group(0)
        if token.startswith("http"):
            cleaned = token.rstrip(".,;)")
            trailing = token[len(cleaned):]
            _add_hyperlink(paragraph, cleaned, cleaned, font_name, size)
            if trailing:
                _set_run_font(paragraph.add_run(trailing), font_name, size)
        elif token.startswith("**"):
            _set_run_font(paragraph.add_run(token[2:-2]), font_name, size, bold=True)
        elif token.startswith("*"):
            _set_run_font(paragraph.add_run(token[1:-1]), font_name, size, italic=True)
        else:
            _set_run_font(paragraph.add_run(token[1:-1]), "Consolas", max(8, size - 0.5), color=INK)
        cursor = match.end()
    if cursor < len(text):
        _set_run_font(paragraph.add_run(text[cursor:]), font_name, size)


def _font_paths(chinese: bool):
    fonts = Path("C:/Windows/Fonts")
    if chinese:
        regular = fonts / "msyh.ttc"
        bold = fonts / "msyhbd.ttc"
    else:
        regular = fonts / "arial.ttf"
        bold = fonts / "arialbd.ttf"
    return regular, bold


def _image_fonts(chinese: bool):
    regular, bold = _font_paths(chinese)
    return {
        "title": ImageFont.truetype(str(bold), 38),
        "head": ImageFont.truetype(str(bold), 25),
        "body": ImageFont.truetype(str(regular), 20),
        "small": ImageFont.truetype(str(regular), 17),
    }


def _draw_centered(draw, box, text, font, fill, spacing=5) -> None:
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=spacing)
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - w) / 2, (y1 + y2 - h) / 2), text,
                        font=font, fill=fill, align="center", spacing=spacing)


def _rounded(draw, box, fill, outline, title, body, fonts) -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=4)
    x1, y1, x2, y2 = box
    _draw_centered(draw, (x1 + 12, y1 + 8, x2 - 12, y1 + 67), title, fonts["head"], "#20354A")
    _draw_centered(draw, (x1 + 16, y1 + 67, x2 - 16, y2 - 10), body, fonts["small"], "#40566A")


def _arrow(draw, start, end, color="#6A7A89", width=4) -> None:
    draw.line([start, end], fill=color, width=width)
    x, y = end
    draw.polygon([(x, y), (x - 14, y - 8), (x - 14, y + 8)], fill=color)


def draw_authority_figure(path: Path, chinese: bool) -> None:
    img = Image.new("RGB", (1900, 1160), "white")
    d = ImageDraw.Draw(img)
    f = _image_fonts(chinese)
    title = "OpenPBL 教育 AI 权威拓扑" if chinese else "OpenPBL Authority Topology for Educational AI"
    subtitle = ("直接帮助的教育意义取决于范围、可见性、可撤销性与高后果权力"
                if chinese else "Educational value depends on scope, visibility, reversibility, and retained consequence")
    _draw_centered(d, (80, 25, 1820, 90), title, f["title"], "#20354A")
    _draw_centered(d, (120, 92, 1780, 132), subtitle, f["small"], "#677788")

    actors = (["教师\n课程审查 · 最终裁决", "学生\n发起 · 核验 · 提交", "AI/系统\n生成 · 辅导 · 受限执行"]
              if chinese else ["Teacher\nCurriculum oversight · final adjudication", "Student\nInitiate · verify · submit", "AI / system\nGenerate · coach · bounded action"])
    actor_boxes = [(95, 170, 570, 310), (712, 170, 1187, 310), (1330, 170, 1805, 310)]
    actor_colors = [(PALE_BLUE, BLUE), (PALE_TEAL, TEAL), (PALE_GOLD, GOLD)]
    for box, text, colors in zip(actor_boxes, actors, actor_colors):
        _rounded(d, box, "#" + colors[0], "#" + colors[1], text.split("\n")[0], text.split("\n")[1], f)

    auth_titles = (["课程权威", "教学权威", "适应性权威", "生产权威", "评价权威"] if chinese
                   else ["Curricular", "Instructional", "Adaptive", "Productive", "Evaluative"])
    auth_bodies = (["先修 / 新授 / 拓展", "解释→练习→终结测评", "缺口补救→完整主课", "明确委托→记录→撤销", "证据门控→教师判断"] if chinese
                   else ["Prior / lesson / extension", "Explain → practice\n→ terminal assessment", "Gap repair → full lesson", "Explicit task → provenance\n→ undo", "Evidence gate\n→ teacher decision"])
    fills = [PALE_BLUE, PALE_TEAL, PALE_GOLD, PALE_GRAY, PALE_RED]
    outlines = [BLUE, TEAL, GOLD, MUTED, "B55345"]
    x0, gap, width = 70, 24, 333
    boxes = []
    for i in range(5):
        box = (x0 + i * (width + gap), 425, x0 + i * (width + gap) + width, 670)
        boxes.append(box)
        _rounded(d, box, "#" + fills[i], "#" + outlines[i], auth_titles[i], auth_bodies[i], f)

    for x in (332, 950, 1568):
        d.line([(x, 310), (x, 375)], fill="#8A98A5", width=4)
        d.polygon([(x, 392), (x - 9, 375), (x + 9, 375)], fill="#8A98A5")
    d.line([(160, 385), (1740, 385)], fill="#C2CBD2", width=3)
    for box in boxes:
        cx = (box[0] + box[2]) // 2
        d.line([(cx, 385), (cx, 414)], fill="#8A98A5", width=3)

    safeguard_title = "可执行护栏" if chinese else "Executable safeguards"
    safeguard = ("类型化课程图 · 教学顺序策略 · 白名单操作 · 冲突保护 · 证据结构 · 人类最终决定"
                 if chinese else "Typed curriculum graph · sequencing policy · whitelisted operations · conflict protection · evidence schema · human final decision")
    d.rounded_rectangle((160, 790, 1740, 920), radius=22, fill="#20354A", outline="#20354A")
    _draw_centered(d, (190, 802, 510, 908), safeguard_title, f["head"], "#FFFFFF")
    _draw_centered(d, (520, 802, 1710, 908), safeguard, f["body"], "#FFFFFF")
    for box in boxes:
        cx = (box[0] + box[2]) // 2
        d.line([(cx, 670), (cx, 756)], fill="#8A98A5", width=3)
        d.polygon([(cx, 775), (cx - 9, 757), (cx + 9, 757)], fill="#8A98A5")
    footer = ("不可委托：AI 提交作品、完成阶段、确认教师决定或作最终评价"
              if chinese else "Non-delegable: AI submission, stage completion, teacher confirmation, or final evaluation")
    _draw_centered(d, (120, 980, 1780, 1045), footer, f["head"], "#A23D31")
    img.save(path, dpi=(180, 180))


def draw_sequence_figure(path: Path, chinese: bool) -> None:
    img = Image.new("RGB", (1900, 1050), "white")
    d = ImageDraw.Draw(img)
    f = _image_fonts(chinese)
    title = "从课程边界到项目证据的人机协作序列" if chinese else "Human-AI Collaboration from Curriculum Boundary to Project Evidence"
    _draw_centered(d, (70, 20, 1830, 90), title, f["title"], "#20354A")

    top_titles = (["角色感知\n课程图", "先修诊断", "缺口专项补救", "完整主课", "唯一终结性\n掌握测评", "有余时拓展"] if chinese else
                  ["Role-aware\ncurriculum graph", "Prerequisite\ndiagnostic", "Gap-specific\nremediation", "Full main lesson", "Single terminal\nmastery assessment", "Enrichment\nif time remains"])
    top_bodies = (["先修≠新授", "一知识点一题", "一缺口一资源", "解释·示例·练习·反馈", "4–8题 · 按知识点", "无评分 · 无后测"] if chinese else
                  ["prior ≠ lesson", "one item per point", "one resource per gap", "explain · example\npractice · feedback", "4–8 tagged items", "ungraded · no post-test"])
    top_fills = [PALE_BLUE, PALE_GOLD, PALE_GOLD, PALE_TEAL, PALE_BLUE, PALE_GRAY]
    top_out = [BLUE, GOLD, GOLD, TEAL, BLUE, MUTED]
    x0, w, gap = 42, 280, 28
    for i in range(6):
        box = (x0 + i * (w + gap), 145, x0 + i * (w + gap) + w, 385)
        title_lines = top_titles[i].split("\n", 1)
        title_text = title_lines[0] if len(title_lines) == 1 else "\n".join(title_lines)
        _rounded(d, box, "#" + top_fills[i], "#" + top_out[i], title_text, top_bodies[i], f)
        if i < 5:
            _arrow(d, (box[2] + 3, 265), (box[2] + gap - 5, 265))

    turn = "所有学生返回完整主课" if chinese else "All learners return to the full main lesson"
    _draw_centered(d, (550, 405, 1350, 455), turn, f["small"], "#177E89")

    lower_titles = (["真实项目任务", "明确委托 AI", "白名单草稿修改", "测试—解释—修订", "学生最终提交", "教师裁决"] if chinese else
                    ["Authentic\nproject", "Explicit AI\ndelegation", "Whitelisted\ndraft edit", "Test–interpret\n–revise", "Student final\nsubmission", "Teacher\nadjudication"])
    lower_bodies = (["问题·方案·制作", "任务与目标可见", "前后值·来源·撤销", "迭代证据链", "认可最终作品", "准备度·评价·干预"] if chinese else
                    ["problem · proposal\n· making", "visible task\nand target", "before/after · source\n· undo", "linked iteration\nevidence", "endorses final\nartifact", "readiness · evaluation\n· intervention"])
    for i in range(6):
        box = (x0 + i * (w + gap), 565, x0 + i * (w + gap) + w, 805)
        fill = [PALE_TEAL, PALE_GOLD, PALE_GRAY, PALE_BLUE, PALE_TEAL, PALE_RED][i]
        out = [TEAL, GOLD, MUTED, BLUE, TEAL, "B55345"][i]
        _rounded(d, box, "#" + fill, "#" + out, lower_titles[i], lower_bodies[i], f)
        if i < 5:
            _arrow(d, (box[2] + 3, 685), (box[2] + gap - 5, 685))

    d.line([(1768, 385), (1810, 385), (1810, 505), (42, 505), (42, 565)], fill="#8A98A5", width=4)
    d.polygon([(42, 565), (34, 549), (50, 549)], fill="#8A98A5")
    boundary = ("系统可建议、执行受限操作并提示风险；学生拥有提交权，教师拥有高后果评价权"
                if chinese else "The system may recommend, execute bounded operations, and surface risk; students submit and teachers retain consequential judgment")
    _draw_centered(d, (110, 880, 1790, 955), boundary, f["head"], "#20354A")
    img.save(path, dpi=(180, 180))


def configure(doc: Document, chinese: bool) -> tuple[str, float]:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.86)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    font_name = "SimSun" if chinese else "Times New Roman"
    size = 10.5 if chinese else 10.8
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.font.size = Pt(size)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    heading_sizes = {1: 15.5, 2: 12.8, 3: 11.2}
    for level, hsize in heading_sizes.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(hsize)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK if level == 1 else BLUE)
        style.paragraph_format.space_before = Pt(11 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    running = "OPENPBL | EDUCATIONAL DESIGN RESEARCH" if not chinese else "OPENPBL｜教育设计研究稿"
    _set_run_font(header.add_run(running), font_name, 7.8, bold=True, color=MUTED)
    _page_number(section.footer.paragraphs[0])
    for run in section.footer.paragraphs[0].runs:
        _set_run_font(run, font_name, 8, color=MUTED)
    return font_name, size


def add_title_block(doc: Document, lines: list[str], chinese: bool, font_name: str) -> int:
    title = lines[0][2:].strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    _set_run_font(p.add_run(title), font_name, 18.5 if chinese else 19.5, bold=True, color=INK)

    consumed = 1
    while consumed < len(lines) and (not lines[consumed].strip() or lines[consumed].startswith("**")):
        line = lines[consumed].strip()
        if line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            _inline(p, line.replace("  ", ""), font_name, 9.2)
        consumed += 1
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(7)
    rule.paragraph_format.space_after = Pt(7)
    ppr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    pbdr.append(bottom)
    ppr.append(pbdr)
    return consumed


def add_table(doc: Document, rows: list[list[str]], font_name: str, font_size: float) -> None:
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    widths = [Inches(6.70 / cols)] * cols
    for r_idx, row in enumerate(rows):
        _keep_row_together(table.rows[r_idx])
        if r_idx == 0:
            _repeat_header(table.rows[r_idx])
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.width = widths[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _cell_margins(cell)
            if r_idx == 0:
                _shade(cell, INK)
            elif r_idx % 2 == 0:
                _shade(cell, PALE_GRAY)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.keep_with_next = r_idx == 0
            _inline(p, value, font_name, max(8.1, font_size - 1.2))
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        raw.append(lines[i].strip())
        i += 1
    rows = []
    for idx, line in enumerate(raw):
        values = [item.strip() for item in line.strip("|").split("|")]
        if idx == 1 and all(re.fullmatch(r":?-{3,}:?", v) for v in values):
            continue
        rows.append(values)
    return rows, i


def insert_figure(doc: Document, image_path: Path, caption: str, font_name: str, chinese: bool) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(image_path), width=Inches(6.55))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    cap.paragraph_format.space_after = Pt(7)
    _set_run_font(cap.add_run(caption), font_name, 8.8, italic=True, color=MUTED)


def build(markdown: Path, output: Path, chinese: bool) -> None:
    lines = markdown.read_text(encoding="utf-8").splitlines()
    doc = Document()
    font_name, body_size = configure(doc, chinese)
    i = add_title_block(doc, lines, chinese, font_name)
    in_references = False
    paragraph_buffer: list[str] = []

    authority_image = ASSETS / ("authority_topology_zh.png" if chinese else "authority_topology_en.png")
    sequence_image = ASSETS / ("learning_sequence_zh.png" if chinese else "learning_sequence_en.png")

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        paragraph_buffer = []
        if not text:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.22
        if in_references:
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.keep_together = True
            _inline(p, text, font_name, 9.2)
        else:
            _inline(p, text, font_name, body_size)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            i += 1
            continue
        if stripped == "{{FIGURE_COAGENCY}}":
            flush_paragraph()
            caption = ("图 1. OpenPBL 的五维权威拓扑及其可执行护栏。" if chinese
                       else "Figure 1. OpenPBL's five-part authority topology and executable safeguards.")
            insert_figure(doc, authority_image, caption, font_name, chinese)
            i += 1
            continue
        if stripped == "{{FIGURE_SEQUENCE}}":
            flush_paragraph()
            caption = ("图 2. 从课程边界、适应性教学到证据门控项目协作的完整序列。" if chinese
                       else "Figure 2. Sequence from curricular boundaries and adaptive teaching to evidence-gated project collaboration.")
            insert_figure(doc, sequence_image, caption, font_name, chinese)
            i += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            rows, i = parse_table(lines, i)
            add_table(doc, rows, font_name, body_size)
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 2")
            _inline(p, stripped[4:], font_name, 12.8)
            i += 1
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            heading = stripped[3:]
            if heading in ("10. Conclusion", "10. 结论"):
                doc.add_page_break()
            if heading in ("References", "参考文献"):
                doc.add_page_break()
                in_references = True
            p = doc.add_paragraph(style="Heading 1")
            _inline(p, heading, font_name, 15.5)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.22)
            p.paragraph_format.space_after = Pt(3)
            _inline(p, stripped, font_name, body_size)
            i += 1
            continue
        if stripped.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(3)
            _inline(p, stripped[2:], font_name, body_size)
            i += 1
            continue
        paragraph_buffer.append(stripped)
        i += 1
    flush_paragraph()

    # Keep section headings with following text and apply journal-review metadata.
    props = doc.core_properties
    props.title = lines[0][2:].strip()
    props.subject = "Educational design research and formative artifact evaluation"
    props.author = "Anonymous for peer review"
    props.keywords = "AIED; project-based learning; human-AI co-agency; learner agency"
    doc.save(output)


def main() -> None:
    draw_authority_figure(ASSETS / "authority_topology_en.png", chinese=False)
    draw_authority_figure(ASSETS / "authority_topology_zh.png", chinese=True)
    draw_sequence_figure(ASSETS / "learning_sequence_en.png", chinese=False)
    draw_sequence_figure(ASSETS / "learning_sequence_zh.png", chinese=True)
    build(
        ROOT / "openpbl_education_manuscript_en_v2.md",
        OUT / "OpenPBL_Education_Journal_Manuscript_EN_FINAL.docx",
        chinese=False,
    )
    build(
        ROOT / "openpbl_education_manuscript_zh_v2.md",
        OUT / "OpenPBL_Education_Journal_Manuscript_ZH_FINAL.docx",
        chinese=True,
    )


if __name__ == "__main__":
    main()
