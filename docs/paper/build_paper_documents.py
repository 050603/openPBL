from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables"
ASSETS = ROOT / "generated_assets"
OUT.mkdir(parents=True, exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)

NAVY = "17324D"
BLUE = "2166AC"
TEAL = "1B7F79"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E8F4F2"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "667788"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def font_for_run(run, name: str, size: float | None = None, bold=None, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline_markdown(paragraph, text: str, font_name: str, size: float) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            font_for_run(paragraph.add_run(text[cursor:match.start()]), font_name, size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            font_for_run(run, font_name, size, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            font_for_run(run, font_name, size)
            run.italic = True
        else:
            run = paragraph.add_run(token[1:-1])
            font_for_run(run, "Consolas", size - 0.5, color=NAVY)
        cursor = match.end()
    if cursor < len(text):
        font_for_run(paragraph.add_run(text[cursor:]), font_name, size)


def configure_document(doc: Document, chinese: bool = False) -> None:
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.15)
    sec.right_margin = Cm(2.15)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    base_font = "Microsoft YaHei" if chinese else "Times New Roman"
    normal = doc.styles["Normal"]
    normal.font.name = base_font
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), base_font)
    normal.font.size = Pt(10.5 if not chinese else 10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    for idx, size in ((1, 15), (2, 12.5), (3, 11)):
        style = doc.styles[f"Heading {idx}"]
        style.font.name = base_font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), base_font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if idx == 1 else BLUE)
        style.paragraph_format.space_before = Pt(10 if idx == 1 else 7)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    text = "OPENPBL RESEARCH MANUSCRIPT" if not chinese else "openPBL 工程研究转化路线图"
    font_for_run(header.add_run(text), base_font, 8, bold=True, color=MID_GRAY)
    add_page_number(sec.footer.paragraphs[0])


def _fonts():
    regular_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    return {
        "title": ImageFont.truetype(str(bold_path), 34),
        "head": ImageFont.truetype(str(bold_path), 22),
        "body": ImageFont.truetype(str(regular_path), 18),
        "small": ImageFont.truetype(str(regular_path), 16),
    }


def _rounded_box(draw, rect, fill, outline, title, body="", fonts=None):
    fonts = fonts or _fonts()
    draw.rounded_rectangle(rect, radius=14, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = rect
    center = (x1 + x2) / 2
    title_box = draw.multiline_textbbox((0, 0), title, font=fonts["head"], align="center", spacing=3)
    title_w = title_box[2] - title_box[0]
    draw.multiline_text((center - title_w / 2, y1 + 18), title, font=fonts["head"], fill="#17324D",
                        align="center", spacing=3)
    if body:
        body_box = draw.multiline_textbbox((0, 0), body, font=fonts["body"], align="center", spacing=4)
        body_w = body_box[2] - body_box[0]
        draw.multiline_text((center - body_w / 2, y1 + 64), body, font=fonts["body"], fill="#34495E",
                            align="center", spacing=4)


def _arrow(draw, start, end, color="#667788", width=4):
    draw.line([start, end], fill=color, width=width)
    x, y = end
    draw.polygon([(x, y), (x - 12, y - 7), (x - 12, y + 7)], fill=color)


def draw_architecture(path: Path) -> None:
    img = Image.new("RGB", (1944, 1044), "white")
    draw = ImageDraw.Draw(img)
    fonts = _fonts()
    title = "OpenPBL Evidence-Gated Orchestration"
    bbox = draw.textbbox((0, 0), title, font=fonts["title"])
    draw.text(((1944 - (bbox[2] - bbox[0])) / 2, 35), title, font=fonts["title"], fill="#17324D")

    top = [
        ((60, 140, 600, 285), "#EAF2F8", "#2166AC", "Teacher surface", "approval • directives • intervention"),
        ((702, 140, 1242, 285), "#E8F4F2", "#1B7F79", "Student project surface", "choose • make • verify • reflect"),
        ((1344, 140, 1884, 285), "#F4ECF7", "#7D3C98", "Role-delimited companions", "explain • question • critique • review"),
    ]
    middle = [
        ((55, 380, 445, 520), "#EAF2F8", "#2166AC", "Stage policies", "role + action boundaries"),
        ((535, 380, 925, 520), "#E8F4F2", "#1B7F79", "Evidence engine", "scope + content location"),
        ((1015, 380, 1405, 520), "#FDF2E9", "#CA6F1E", "Orchestrator", "cooldown + speaker budget"),
        ((1495, 380, 1885, 520), "#FDEDEC", "#B03A2E", "Stage gates", "block • warn • override"),
    ]
    evidence = [
        ((55, 640, 445, 780), "#F3F5F7", "#667788", "Typed course events", "46 action variants"),
        ((535, 640, 925, 780), "#F3F5F7", "#667788", "Learning evidence", "artifacts • checks • adoption"),
        ((1015, 640, 1405, 780), "#F3F5F7", "#667788", "Human decisions", "student + teacher confirmation"),
        ((1495, 640, 1885, 780), "#F3F5F7", "#667788", "Persistent audit", "44 data models"),
    ]
    bottom = [
        ((135, 880, 595, 990), "#F8F9F9", "#566573", "PostgreSQL / Prisma", ""),
        ((742, 880, 1202, 990), "#F8F9F9", "#566573", "WebSocket + polling fallback", ""),
        ((1349, 880, 1809, 990), "#F8F9F9", "#566573", "Provider-agnostic LLM + TTS", ""),
    ]
    for item in top + middle + evidence + bottom:
        _rounded_box(draw, *item, fonts=fonts)
    for x in (330, 972, 1614):
        draw.line([(x, 285), (x, 350)], fill="#667788", width=4)
        draw.polygon([(x, 364), (x - 8, 350), (x + 8, 350)], fill="#667788")
    for x in (250, 730, 1210, 1690):
        draw.line([(x, 520), (x, 610)], fill="#667788", width=4)
        draw.polygon([(x, 624), (x - 8, 610), (x + 8, 610)], fill="#667788")
    loop_text = "evidence → scoped inference → learner action or teacher decision → new evidence"
    loop_box = draw.textbbox((0, 0), loop_text, font=fonts["small"])
    draw.text(((1944 - (loop_box[2] - loop_box[0])) / 2, 560), loop_text,
              font=fonts["small"], fill="#2166AC")
    img.save(path)


def draw_research_program(path: Path) -> None:
    img = Image.new("RGB", (1944, 864), "white")
    draw = ImageDraw.Draw(img)
    fonts = _fonts()
    title = "Engineering Artifact to Research Program"
    bbox = draw.textbbox((0, 0), title, font=fonts["title"])
    draw.text(((1944 - (bbox[2] - bbox[0])) / 2, 55), title, font=fonts["title"], fill="#17324D")
    stages = [
        ((55, 230, 430, 515), "#EAF2F8", "#2166AC", "1. Fidelity", "Design principles\nCore tests\nTraceability"),
        ((540, 230, 915, 515), "#E8F4F2", "#1B7F79", "2. Feasibility", "Teacher co-design\nUsability\nPilot reliability"),
        ((1025, 230, 1400, 515), "#FDF2E9", "#CA6F1E", "3. Effectiveness", "Cluster trial\nTransfer + agency\nTeacher workload"),
        ((1510, 230, 1885, 515), "#F4ECF7", "#7D3C98", "4. Generalization", "Ablations\nMultiple schools\nOpen materials"),
    ]
    for item in stages:
        _rounded_box(draw, *item, fonts=fonts)
    for start, end in (((430, 372), (525, 372)), ((915, 372), (1010, 372)), ((1400, 372), (1495, 372))):
        _arrow(draw, start, end)
    warning = "No learning-effect claim before Stage 3"
    wb = draw.textbbox((0, 0), warning, font=fonts["head"])
    draw.text(((1944 - (wb[2] - wb[0])) / 2, 645), warning, font=fonts["head"], fill="#B03A2E")
    note = "Freeze code, prompts, model versions, measures, and analysis plan"
    nb = draw.textbbox((0, 0), note, font=fonts["body"])
    draw.text(((1944 - (nb[2] - nb[0])) / 2, 700), note, font=fonts["body"], fill="#667788")
    img.save(path)


def add_table(doc: Document, rows: list[list[str]], font_name: str, font_size: float) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_markdown(p, text, font_name, font_size)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_markdown_to_docx(md_path: Path, out_path: Path, chinese: bool = False,
                           architecture: Path | None = None, program: Path | None = None) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc, chinese=chinese)
    font_name = "Microsoft YaHei" if chinese else "Times New Roman"
    body_size = 10 if chinese else 10.3
    i = 0
    title_done = False

    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1]):
            table_lines = [stripped]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in row.strip("|").split("|")] for row in table_lines]
            add_table(doc, rows, font_name, 8.4 if not chinese else 8.2)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:])
            font_for_run(run, font_name, 20 if chinese else 21, bold=True, color=NAVY)
            title_done = True
            i += 1
            continue
        if stripped.startswith("## "):
            heading_text = stripped[3:]
            doc.add_heading(heading_text, level=1)
            if architecture and heading_text == "5. System Architecture and Implementation":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(architecture), width=Inches(6.45))
                cap = doc.add_paragraph("Figure 1. OpenPBL evidence-gated system architecture.")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                font_for_run(cap.runs[0], font_name, 8.5, color=MID_GRAY)
            if program and heading_text in ("四、真实课堂研究的推荐设计", "8. Classroom Evaluation Protocol"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(program), width=Inches(6.45))
                cap_text = "图 1. 从工程原型到可验证研究成果的递进路径。" if chinese else \
                    "Figure 2. Progression from an engineering artifact to a generalizable research program."
                cap = doc.add_paragraph(cap_text)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                font_for_run(cap.runs[0], font_name, 8.5, color=MID_GRAY)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.right_indent = Cm(0.55)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            add_inline_markdown(p, stripped[2:], font_name, body_size)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor.from_string(TEAL)
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, re.sub(r"^[-*]\s+", "", stripped), font_name, body_size)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            add_inline_markdown(p, stripped, font_name, body_size)
            i += 1
            continue

        paragraph_parts = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith("#") or nxt.startswith("|") or nxt.startswith("> ")
                    or re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt)):
                break
            paragraph_parts.append(nxt)
            i += 1
        text = " ".join(paragraph_parts)
        p = doc.add_paragraph()
        if not title_done:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_markdown(p, text, font_name, body_size)
        if text.startswith("**Anonymous Author"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
        if text.startswith("**Keywords:"):
            p.paragraph_format.space_after = Pt(8)
        if re.match(r"^\[\d+\]", text):
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9)

    doc.core_properties.title = lines[0].lstrip("# ").strip()
    doc.core_properties.subject = "OpenPBL research manuscript" if not chinese else "openPBL 工程研究转化"
    doc.core_properties.author = "Anonymous Author(s)" if not chinese else "openPBL Research Team"
    doc.core_properties.keywords = "OpenPBL, project-based learning, generative AI, evidence-gated orchestration"
    doc.save(out_path)


def main() -> None:
    arch = ASSETS / "openpbl_architecture.png"
    program = ASSETS / "openpbl_research_program.png"
    draw_architecture(arch)
    draw_research_program(program)
    parse_markdown_to_docx(
        ROOT / "openpbl_manuscript.md",
        OUT / "OpenPBL_English_Manuscript.docx",
        chinese=False,
        architecture=arch,
        program=program,
    )
    parse_markdown_to_docx(
        ROOT / "openpbl_engineering_to_research_roadmap_cn.md",
        OUT / "OpenPBL_Engineering_to_Research_Roadmap_CN.docx",
        chinese=True,
        program=program,
    )
    print(OUT / "OpenPBL_English_Manuscript.docx")
    print(OUT / "OpenPBL_Engineering_to_Research_Roadmap_CN.docx")


if __name__ == "__main__":
    main()
