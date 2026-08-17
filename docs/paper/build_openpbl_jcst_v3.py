from __future__ import annotations

import json
import math
import re
from dataclasses import dataclass
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_bilingual_education_paper as legacy


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables"
ASSETS_V2 = ROOT / "generated_assets_v2"
ASSETS_V3 = ROOT / "generated_assets_v3"
DATA_DIR = ROOT / "experimental_placeholders"
EVIDENCE_DIR = ROOT / "research_evidence_2026-08-13_v3"
for directory in (OUT, ASSETS_V3, DATA_DIR, EVIDENCE_DIR):
    directory.mkdir(parents=True, exist_ok=True)

EN_SOURCE = ROOT / "openpbl_education_manuscript_en_v2.md"
ZH_SOURCE = ROOT / "openpbl_education_manuscript_zh_v2.md"
EN_V3 = ROOT / "openpbl_jcst_manuscript_en_v3.md"
ZH_V3 = ROOT / "openpbl_jcst_manuscript_zh_v3.md"
EN_DOCX = OUT / "OpenPBL_JCST_Style_Manuscript_EN_v3.docx"
BI_DOCX = OUT / "OpenPBL_JCST_Style_Bilingual_Review_v3.docx"

SEED = 20260813
INK = "20354A"
BLUE = "2E74B5"
TEAL = "177E89"
GOLD = "B7791F"
MUTED = "677788"
PALE_BLUE = "EAF2F8"
PALE_GRAY = "F2F5F7"
PALE_RED = "FCEBE8"
WHITE = "FFFFFF"
BLACK = "111111"
SYNTH_RED = "A23D31"


REFERENCES = [
    "[1] Kestin G, Miller K, Klales A, Milbourne T, Ponti G. AI tutoring outperforms in-class active learning: An RCT introducing a novel research-based design in an authentic educational setting. *Scientific Reports*, 2025, 15: 17458. https://doi.org/10.1038/s41598-025-97652-6",
    "[2] Holstein K, McLaren B M, Aleven V. Co-designing a real-time classroom orchestration tool to support teacher-AI complementarity. *Journal of Learning Analytics*, 2019, 6(2): 27-52. https://doi.org/10.18608/jla.2019.62.3",
    "[3] Bastani H, Bastani O, Sungu A, Ge H, Kabakcı Ö, Mariman R. Generative AI without guardrails can harm learning: Evidence from high school mathematics. *Proceedings of the National Academy of Sciences*, 2025, 122(26): e2422633122. https://doi.org/10.1073/pnas.2422633122",
    "[4] Darvishi A, Khosravi H, Sadiq S, Gašević D, Siemens G. Impact of AI assistance on student agency. *Computers & Education*, 2024, 210: 104967. https://doi.org/10.1016/j.compedu.2023.104967",
    "[5] Chen C-H, Yang Y-C. Revisiting the effects of project-based learning on students' academic achievement: A meta-analysis investigating moderators. *Educational Research Review*, 2019, 26: 71-81. https://doi.org/10.1016/j.edurev.2018.11.001",
    "[6] Farshad S, Fortin C. An umbrella review of meta-analyses on project-based learning: Effects on academic achievement, higher-order thinking, and 21st-century skills. *Educational Research Review*, 2026, 52: 100809. https://doi.org/10.1016/j.edurev.2026.100809",
    "[7] Hmelo-Silver C E, Duncan R G, Chinn C A. Scaffolding and achievement in problem-based and inquiry learning: A response to Kirschner, Sweller, and Clark (2006). *Educational Psychologist*, 2007, 42(2): 99-107. https://doi.org/10.1080/00461520701263368",
    "[8] Kim N J, Belland B R, Walker A E. Effectiveness of computer-based scaffolding in the context of problem-based learning for STEM education: Bayesian meta-analysis. *Educational Psychology Review*, 2018, 30: 397-429. https://doi.org/10.1007/s10648-017-9419-1",
    "[9] Kim J, Lee H, Cho Y H. Learning design to support student-AI collaboration: Perspectives of leading teachers for AI in education. *Education and Information Technologies*, 2022, 27: 6069-6104. https://doi.org/10.1007/s10639-021-10831-6",
    "[10] Yu J F, Zhang-Li D, Zhang Z Y, et al. From MOOC to MAIC: Reimagine online teaching and learning through LLM-driven agents. *Journal of Computer Science and Technology*, 2026, 41(1): 394-414. https://doi.org/10.1007/s11390-025-6000-0",
    "[11] Samuel A. Learning with machines: Toward a theory of epistemic co-agency. *Computers and Education: Artificial Intelligence*, 2026, 10: 100573. https://doi.org/10.1016/j.caeai.2026.100573",
    "[12] Chaaban Y, Jung S-G, Medina J, Azem J Y, Salminen J, Jansen B J. Examining student teachers' agency in an AI-supported learning environment: Q methodology research. *International Journal of Artificial Intelligence in Education*, 2025, 35: 4083-4107. https://doi.org/10.1007/s40593-025-00529-y",
    "[13] Kim J, Detrick R, Yu S, Song Y, Bol L, Li N. Socially shared regulation of learning and artificial intelligence: Opportunities to support socially shared regulation. *Education and Information Technologies*, 2025, 30: 11483-11521. https://doi.org/10.1007/s10639-024-13187-9",
    "[14] Mouta A, Torrecilla-Sánchez E M, Pinto-Llorente A M. Where is agency moving to? Exploring the interplay between AI technologies in education and human agency. *Digital Society*, 2025, 4: 49. https://doi.org/10.1007/s44206-025-00203-9",
    "[15] Kasneci E, Sessler K, Küchemann S, et al. ChatGPT for good? On opportunities and challenges of large language models for education. *Learning and Individual Differences*, 2023, 103: 102274. https://doi.org/10.1016/j.lindif.2023.102274",
    "[16] Chanaa A, El Faddouli N. Prerequisites-based course recommendation: Recommending learning objects using concept prerequisites and metadata matching. *Smart Learning Environments*, 2024, 11: 16. https://doi.org/10.1186/s40561-024-00301-0",
    "[17] Hadwin A F, Järvelä S, Miller M. Self-regulation, co-regulation, and shared regulation in collaborative learning environments. In: Schunk D H, Greene J A, eds. *Handbook of Self-Regulation of Learning and Performance*, 2nd ed. Routledge, 2018: 83-106. https://doi.org/10.4324/9781315697048-6",
    "[18] Banihashem S K, Bond M, Bergdahl N, Khosravi H, Noroozi O. A systematic mapping review at the intersection of artificial intelligence and self-regulated learning. *International Journal of Educational Technology in Higher Education*, 2025, 22: 50. https://doi.org/10.1186/s41239-025-00548-8",
    "[19] de Mooij S, Lämsä J, Lim L, et al. A systematic review of self-regulated learning through integration of multimodal data and artificial intelligence. *Educational Psychology Review*, 2025, 37: 54. https://doi.org/10.1007/s10648-025-10028-0",
    "[20] Dillenbourg P. Design for classroom orchestration. *Computers & Education*, 2013, 69: 485-492. https://doi.org/10.1016/j.compedu.2013.04.013",
    "[21] Possaghi I, Vesin B, Zhang F, et al. Integrating multi-modal learning analytics dashboard in K-12 education: Insights for enhancing orchestration and teacher decision-making. *Smart Learning Environments*, 2025, 12: 53. https://doi.org/10.1186/s40561-025-00410-4",
    "[22] Wang F, Hannafin M J. Design-based research and technology-enhanced learning environments. *Educational Technology Research and Development*, 2005, 53: 5-23. https://doi.org/10.1007/BF02504682",
    "[23] Miao F, Holmes W. *Guidance for Generative AI in Education and Research*. Paris: UNESCO, 2023.",
]


CITATIONS = {
    "(Hmelo-Silver et al., 2007; Kim et al., 2018)": "[7, 8]",
    "(Holstein et al., 2019; Possaghi et al., 2025)": "[2, 21]",
    "(Kestin et al., 2025)": "[1]",
    "(Holstein et al., 2019)": "[2]",
    "(Bastani et al., 2025)": "[3]",
    "(Darvishi et al., 2024)": "[4]",
    "(Chen & Yang, 2019)": "[5]",
    "(Farshad & Fortin, 2026)": "[6]",
    "(Hmelo-Silver et al., 2007)": "[7]",
    "(Kim et al., 2018)": "[8]",
    "(Kim, Lee, & Cho, 2022)": "[9]",
    "(Samuel, 2026)": "[11]",
    "(Chaaban et al., 2025)": "[12]",
    "(Kim et al., 2025)": "[13]",
    "(Mouta et al., 2025)": "[14]",
    "(Kasneci et al., 2023)": "[15]",
    "(Chanaa & El Faddouli, 2024)": "[16]",
    "(Hadwin, Järvelä, & Miller, 2018)": "[17]",
    "(Banihashem et al., 2025)": "[18]",
    "(de Mooij et al., 2025)": "[19]",
    "(Dillenbourg, 2013)": "[20]",
    "(Wang & Hannafin, 2005)": "[22]",
    "(Miao & Holmes, 2023)": "[23]",
    "Kestin et al. (2025)": "Kestin et al. [1]",
    "Bastani et al. (2025)": "Bastani et al. [3]",
    "Darvishi et al. (2024)": "Darvishi et al. [4]",
}

ZH_CITATIONS = {
    "（Hmelo-Silver et al., 2007; Kim et al., 2018）": "[7, 8]",
    "（Holstein et al., 2019; Possaghi et al., 2025）": "[2, 21]",
    "（Kestin et al., 2025）": "[1]",
    "（Holstein et al., 2019）": "[2]",
    "（Bastani et al., 2025）": "[3]",
    "（Darvishi et al., 2024）": "[4]",
    "（Chen & Yang, 2019）": "[5]",
    "（Farshad & Fortin, 2026）": "[6]",
    "（Hmelo-Silver et al., 2007）": "[7]",
    "（Kim et al., 2018）": "[8]",
    "（Kim, Lee, & Cho, 2022）": "[9]",
    "（Samuel, 2026）": "[11]",
    "（Chaaban et al., 2025）": "[12]",
    "（Kim et al., 2025）": "[13]",
    "（Mouta et al., 2025）": "[14]",
    "（Kasneci et al., 2023）": "[15]",
    "（Chanaa & El Faddouli, 2024）": "[16]",
    "（Hadwin, Järvelä, & Miller, 2018）": "[17]",
    "（Banihashem et al., 2025）": "[18]",
    "（de Mooij et al., 2025）": "[19]",
    "（Dillenbourg, 2013）": "[20]",
    "（Wang & Hannafin, 2005）": "[22]",
    "（Miao & Holmes, 2023）": "[23]",
    "Kestin 等（2025）": "Kestin 等 [1]",
    "Bastani 等（2025）": "Bastani 等 [3]",
    "Darvishi 等（2024）": "Darvishi 等 [4]",
}


@dataclass
class Token:
    kind: str
    value: object


def generate_synthetic_data() -> tuple[pd.DataFrame, pd.DataFrame]:
    rng = np.random.default_rng(SEED)
    conditions = ["PBL-only", "Unrestricted LLM", "Role-based AI", "OpenPBL full"]
    rows: list[dict[str, object]] = []
    effects = {
        "PBL-only": dict(post=4.5, delayed=2.0, project=0.0, agency=0.00, offload=0.00, integrity=0.0, minutes=0.0),
        "Unrestricted LLM": dict(post=7.5, delayed=0.5, project=8.0, agency=-0.45, offload=0.90, integrity=-8.0, minutes=-4.0),
        "Role-based AI": dict(post=8.5, delayed=4.0, project=10.0, agency=-0.15, offload=0.40, integrity=0.0, minutes=-3.5),
        "OpenPBL full": dict(post=11.5, delayed=9.0, project=12.0, agency=0.25, offload=-0.15, integrity=20.0, minutes=-6.5),
    }
    participant_id = 1
    for condition in conditions:
        fx = effects[condition]
        for _ in range(60):
            ability = rng.normal()
            engagement = rng.normal()
            pretest = np.clip(68 + 8 * ability + rng.normal(0, 3.2), 35, 94)
            posttest = np.clip(pretest + fx["post"] + 1.1 * engagement + rng.normal(0, 4.0), 35, 100)
            delayed = np.clip(pretest + fx["delayed"] + 1.4 * engagement + rng.normal(0, 4.6), 30, 100)
            project = np.clip(70 + 4.0 * ability + 2.0 * engagement + fx["project"] + rng.normal(0, 5.8), 35, 100)
            agency = np.clip(3.80 + 0.16 * engagement + fx["agency"] + rng.normal(0, 0.42), 1, 5)
            offloading = np.clip(2.45 - 0.10 * ability + fx["offload"] + rng.normal(0, 0.48), 1, 5)
            integrity = np.clip(65 + 3.5 * engagement + fx["integrity"] + rng.normal(0, 7.6), 25, 100)
            minutes = np.clip(15.2 - 0.7 * engagement + fx["minutes"] + rng.normal(0, 2.3), 3, 26)
            rows.append({
                "participant_id": f"S{participant_id:03d}",
                "condition": condition,
                "pretest": round(float(pretest), 2),
                "posttest": round(float(posttest), 2),
                "delayed_transfer": round(float(delayed), 2),
                "project_quality": round(float(project), 2),
                "epistemic_agency_1to5": round(float(agency), 2),
                "cognitive_offloading_1to5": round(float(offloading), 2),
                "evidence_integrity_pct": round(float(integrity), 2),
                "teacher_intervention_minutes": round(float(minutes), 2),
                "synthetic_placeholder": True,
                "random_seed": SEED,
            })
            participant_id += 1
    data = pd.DataFrame(rows)
    ordered = pd.CategoricalDtype(conditions, ordered=True)
    data["condition"] = data["condition"].astype(ordered)
    measures = [
        "pretest", "posttest", "delayed_transfer", "project_quality",
        "epistemic_agency_1to5", "cognitive_offloading_1to5",
        "evidence_integrity_pct", "teacher_intervention_minutes",
    ]
    grouped = data.groupby("condition", observed=False)[measures].agg(["mean", "std", "count"])
    records = []
    for condition in conditions:
        row: dict[str, object] = {"condition": condition, "n": int(grouped.loc[condition, ("pretest", "count")])}
        for measure in measures:
            row[f"{measure}_mean"] = float(grouped.loc[condition, (measure, "mean")])
            row[f"{measure}_sd"] = float(grouped.loc[condition, (measure, "std")])
        records.append(row)
    summary = pd.DataFrame(records)
    data.to_csv(DATA_DIR / f"OpenPBL_synthetic_placeholder_seed{SEED}.csv", index=False, encoding="utf-8-sig")
    summary.to_csv(DATA_DIR / f"OpenPBL_synthetic_summary_seed{SEED}.csv", index=False, encoding="utf-8-sig")
    return data, summary


def _chart_fonts(chinese: bool) -> dict[str, ImageFont.FreeTypeFont]:
    fonts = Path("C:/Windows/Fonts")
    regular = fonts / ("msyh.ttc" if chinese else "arial.ttf")
    bold = fonts / ("msyhbd.ttc" if chinese else "arialbd.ttf")
    return {
        "banner": ImageFont.truetype(str(bold), 31),
        "title": ImageFont.truetype(str(bold), 28),
        "axis": ImageFont.truetype(str(regular), 19),
        "value": ImageFont.truetype(str(bold), 19),
        "small": ImageFont.truetype(str(regular), 16),
    }


def draw_synthetic_figure(summary: pd.DataFrame, path: Path, chinese: bool) -> None:
    image = Image.new("RGB", (1800, 1160), "white")
    draw = ImageDraw.Draw(image)
    fonts = _chart_fonts(chinese)
    draw.rectangle((0, 0, 1800, 82), fill="#A23D31")
    banner = "合成占位数据——不是实证结果，不得用于效果声明" if chinese else "SYNTHETIC PLACEHOLDER — NOT EMPIRICAL RESULTS"
    bbox = draw.textbbox((0, 0), banner, font=fonts["banner"])
    draw.text(((1800 - (bbox[2] - bbox[0])) / 2, 22), banner, fill="white", font=fonts["banner"])
    title = "四种教学条件的预设分析流程输出" if chinese else "Dry-run outputs for the preregistered four-condition analysis"
    bbox = draw.textbbox((0, 0), title, font=fonts["title"])
    draw.text(((1800 - (bbox[2] - bbox[0])) / 2, 100), title, fill="#20354A", font=fonts["title"])

    panels = [
        ("delayed_transfer", "延迟独立迁移" if chinese else "Delayed unaided transfer", (55, 90)),
        ("epistemic_agency_1to5", "认识性能动性（1–5）" if chinese else "Epistemic agency (1–5)", (2.5, 4.6)),
        ("evidence_integrity_pct", "证据链完整度（%）" if chinese else "Evidence-chain integrity (%)", (45, 95)),
        ("teacher_intervention_minutes", "教师干预分钟数（越低越好）" if chinese else "Teacher intervention minutes (lower is better)", (5, 18)),
    ]
    condition_labels = (["PjBL", "无限制\nLLM", "角色型\nAI", "OpenPBL"] if chinese
                        else ["PjBL", "Unrestricted\nLLM", "Role-based\nAI", "OpenPBL"])
    colors = ["#8997A5", "#D69E2E", "#4F86C6", "#177E89"]
    origins = [(85, 195), (925, 195), (85, 660), (925, 660)]
    panel_w, panel_h = 790, 410
    for (measure, panel_title, ylim), (ox, oy) in zip(panels, origins):
        draw.rounded_rectangle((ox, oy, ox + panel_w, oy + panel_h), radius=18, fill="#F8FAFB", outline="#D6DEE4", width=3)
        bbox = draw.textbbox((0, 0), panel_title, font=fonts["title"])
        draw.text((ox + (panel_w - (bbox[2] - bbox[0])) / 2, oy + 18), panel_title, fill="#20354A", font=fonts["title"])
        left, top, right, bottom = ox + 85, oy + 85, ox + panel_w - 30, oy + panel_h - 82
        draw.line((left, top, left, bottom), fill="#6A7A89", width=3)
        draw.line((left, bottom, right, bottom), fill="#6A7A89", width=3)
        ymin, ymax = ylim
        for tick in np.linspace(ymin, ymax, 4):
            y = bottom - (tick - ymin) / (ymax - ymin) * (bottom - top)
            draw.line((left, y, right, y), fill="#E1E6EA", width=2)
            label = f"{tick:.1f}" if ymax <= 5 else f"{tick:.0f}"
            draw.text((left - 52, y - 10), label, fill="#677788", font=fonts["small"])
        slot = (right - left) / 4
        for idx, (_, row) in enumerate(summary.iterrows()):
            mean = float(row[f"{measure}_mean"])
            sd = float(row[f"{measure}_sd"])
            n = int(row["n"])
            se95 = 1.96 * sd / math.sqrt(n)
            center = left + slot * (idx + 0.5)
            bar_w = 78
            y_mean = bottom - (mean - ymin) / (ymax - ymin) * (bottom - top)
            draw.rounded_rectangle((center - bar_w / 2, y_mean, center + bar_w / 2, bottom), radius=8, fill=colors[idx])
            y_hi = bottom - (min(ymax, mean + se95) - ymin) / (ymax - ymin) * (bottom - top)
            y_lo = bottom - (max(ymin, mean - se95) - ymin) / (ymax - ymin) * (bottom - top)
            draw.line((center, y_hi, center, y_lo), fill="#20354A", width=3)
            draw.line((center - 12, y_hi, center + 12, y_hi), fill="#20354A", width=3)
            draw.line((center - 12, y_lo, center + 12, y_lo), fill="#20354A", width=3)
            value_label = f"{mean:.2f}" if ymax <= 5 else f"{mean:.1f}"
            bbox = draw.textbbox((0, 0), value_label, font=fonts["value"])
            draw.text((center - (bbox[2] - bbox[0]) / 2, max(top + 2, y_hi - 27)), value_label, fill="#20354A", font=fonts["value"])
            parts = condition_labels[idx].split("\n")
            for line_idx, part in enumerate(parts):
                bbox = draw.textbbox((0, 0), part, font=fonts["small"])
                draw.text((center - (bbox[2] - bbox[0]) / 2, bottom + 10 + line_idx * 18), part, fill="#40566A", font=fonts["small"])
    footer = (f"固定随机种子 {SEED}；每组 n=60；误差线为合成样本均值的 95% 描述性区间。"
              if chinese else f"Fixed seed {SEED}; n=60 per condition; error bars are 95% descriptive intervals for the synthetic sample means.")
    bbox = draw.textbbox((0, 0), footer, font=fonts["small"])
    draw.text(((1800 - (bbox[2] - bbox[0])) / 2, 1120), footer, fill="#677788", font=fonts["small"])
    image.save(path, dpi=(180, 180))


def fmt(summary: pd.DataFrame, condition: str, measure: str) -> str:
    row = summary.loc[summary["condition"] == condition].iloc[0]
    return f"{row[f'{measure}_mean']:.1f} ({row[f'{measure}_sd']:.1f})"


def synthetic_section(summary: pd.DataFrame, chinese: bool) -> str:
    conditions = ["PBL-only", "Unrestricted LLM", "Role-based AI", "OpenPBL full"]
    if not chinese:
        table3 = [
            "### Table 3. SYNTHETIC PLACEHOLDER descriptive learning outcomes, mean (SD)",
            "",
            "| Condition | n | Pretest | Posttest | Delayed transfer | Project quality |",
            "|---|---:|---:|---:|---:|---:|",
        ]
        labels = {"PBL-only": "PjBL only", "Unrestricted LLM": "PjBL + unrestricted LLM", "Role-based AI": "Role-based AI", "OpenPBL full": "Full OpenPBL"}
        for condition in conditions:
            table3.append(f"| {labels[condition]} | 60 | {fmt(summary, condition, 'pretest')} | {fmt(summary, condition, 'posttest')} | {fmt(summary, condition, 'delayed_transfer')} | {fmt(summary, condition, 'project_quality')} |")
        table4 = [
            "### Table 4. SYNTHETIC PLACEHOLDER agency, evidence, and orchestration outcomes, mean (SD)",
            "",
            "| Condition | Agency (1-5) | Offloading (1-5) | Evidence integrity (%) | Teacher minutes |",
            "|---|---:|---:|---:|---:|",
        ]
        for condition in conditions:
            table4.append(f"| {labels[condition]} | {fmt(summary, condition, 'epistemic_agency_1to5')} | {fmt(summary, condition, 'cognitive_offloading_1to5')} | {fmt(summary, condition, 'evidence_integrity_pct')} | {fmt(summary, condition, 'teacher_intervention_minutes')} |")
        return "\n".join([
            "## 7. Comparative Study Protocol and Synthetic Pipeline Dry Run",
            "",
            "**SYNTHETIC PLACEHOLDER—NOT EMPIRICAL RESULTS.** No participant produced the values in this section. They are seeded simulation records used only to test whether the proposed comparisons, tables, graphics, and analysis code behave coherently. Every value must be replaced with approved human-participant data before submission; none supports a claim that OpenPBL improves learning, agency, or workload.",
            "",
            "### 7.1 Falsifiable comparative design",
            "",
            "The intended study is a preregistered, cluster-randomized comparison with four conditions: (C1) PjBL without generative AI; (C2) PjBL with an unrestricted general-purpose LLM; (C3) role-based AI tutoring and collaboration without OpenPBL's authority compiler, evidence gates, or reversible artifact operations; and (C4) full OpenPBL. C1 estimates the contribution and cost of AI access, C2 tests whether immediate productivity can diverge from independent learning, C3 separates social-role organization from executable authority governance, and C4 tests the complete design. Randomization should occur at class or team level to limit contamination, with baseline achievement and teacher effects handled in the allocation and model.",
            "",
            "The primary confirmatory outcome is delayed, unaided transfer aligned with lesson targets. Secondary outcomes are immediate posttest, artifact quality scored blind to condition, epistemic agency, cognitive offloading, evidence-chain integrity, and teacher intervention minutes per team. The critical contrast is C4 versus C3: it tests whether executable authority boundaries add value beyond role-based AI. C4 versus C2 tests whether guardrails preserve transfer and evidence without erasing productive assistance. Artifact quality must never stand alone because fluent AI production can mask weak independent understanding [3, 4].",
            "",
            "### 7.2 Analysis and reporting plan",
            "",
            "The confirmatory model should estimate delayed transfer from condition while adjusting for pretest, class clustering, and prespecified teacher and subject factors. Secondary models should use the same contrast matrix, multiplicity control, effect sizes with confidence intervals, missing-data sensitivity analyses, and an intention-to-treat estimand. Process analysis should test preregistered mediation by verification and evidence-chain quality, but mediation would remain associational unless its assumptions are defended. Equity analysis should report heterogeneous effects by prior knowledge and accessibility needs without using small subgroups for automated decisions. Adverse outcomes—false prerequisite routing, unverified AI content, fabricated evidence, inaccessible interaction, privacy incidents, and alert fatigue—must be reported alongside benefits.",
            "",
            "### 7.3 Seeded data-generating scenario",
            "",
            f"To debug the pipeline, we generated 240 synthetic records (60 per condition) with NumPy's PCG64 generator and fixed seed {SEED}. Baseline scores were drawn from the same distribution. The scenario deliberately encodes overlapping, imperfect trends rather than deterministic superiority: unrestricted LLM access raises immediate posttest and artifact quality but increases offloading and weakens delayed transfer; role-based AI improves productive outcomes but only partly protects agency; full OpenPBL improves delayed transfer and evidence integrity while reducing teacher intervention time. These are assumptions chosen to expose whether the planned measures can distinguish productivity, learning, agency, and orchestration—not observed effects.",
            "",
            *table3,
            "",
            *table4,
            "",
            "{{FIGURE_SYNTHETIC}}",
            "",
            "### 7.4 Pipeline interpretation",
            "",
            f"In the synthetic dry run, mean delayed transfer was {fmt(summary, 'PBL-only', 'delayed_transfer')} in C1, {fmt(summary, 'Unrestricted LLM', 'delayed_transfer')} in C2, {fmt(summary, 'Role-based AI', 'delayed_transfer')} in C3, and {fmt(summary, 'OpenPBL full', 'delayed_transfer')} in C4. The unrestricted condition produced higher artifact quality than PjBL alone while showing the lowest agency and evidence integrity. Full OpenPBL showed the highest delayed transfer, agency, and evidence integrity and the lowest teacher-intervention time. This joint pattern is useful because it makes the paper's claim falsifiable: real data could show that the controls add friction without learning benefit, that role-based AI is sufficient, or that OpenPBL's advantage is confined to documentation rather than learning.",
            "",
            "No p-value is reported because inferential significance on invented observations would be meaningless. The dry run has three legitimate outputs only: a frozen schema and contrast plan, confirmation that the analysis distinguishes the intended constructs, and an explicit replacement checklist. Before submission, researchers must preregister the protocol, obtain ethics approval, replace the CSV, regenerate Tables 3-4 and Figure 3, archive the analytic code, and rewrite Sections 7.3-7.4 and all claims that depend on them.",
            "",
        ])
    table3 = [
        "### 表 3. 合成占位数据的学习结果描述统计，均值（标准差）",
        "",
        "| 条件 | n | 前测 | 即时后测 | 延迟迁移 | 作品质量 |",
        "|---|---:|---:|---:|---:|---:|",
    ]
    labels = {"PBL-only": "仅 PjBL", "Unrestricted LLM": "PjBL + 无限制 LLM", "Role-based AI": "角色型 AI", "OpenPBL full": "完整 OpenPBL"}
    for condition in conditions:
        table3.append(f"| {labels[condition]} | 60 | {fmt(summary, condition, 'pretest')} | {fmt(summary, condition, 'posttest')} | {fmt(summary, condition, 'delayed_transfer')} | {fmt(summary, condition, 'project_quality')} |")
    table4 = [
        "### 表 4. 合成占位数据的能动性、证据与编排结果，均值（标准差）",
        "",
        "| 条件 | 能动性（1–5） | 认知卸载（1–5） | 证据完整度（%） | 教师分钟数 |",
        "|---|---:|---:|---:|---:|",
    ]
    for condition in conditions:
        table4.append(f"| {labels[condition]} | {fmt(summary, condition, 'epistemic_agency_1to5')} | {fmt(summary, condition, 'cognitive_offloading_1to5')} | {fmt(summary, condition, 'evidence_integrity_pct')} | {fmt(summary, condition, 'teacher_intervention_minutes')} |")
    return "\n".join([
        "## 7. 比较研究方案与合成数据流程试运行",
        "",
        "**合成占位数据——不是实证结果。** 本节数值并非来自任何参与者，而是仅用于检查比较设计、表格、图形和分析代码能否一致运行的固定种子模拟记录。投稿前必须使用经过审批的人类参与者数据替换全部数值；这些数据不能支持 OpenPBL 改善学习、能动性或教师工作量的结论。",
        "",
        "### 7.1 可证伪的比较设计",
        "",
        "拟开展预注册的整群随机研究，设置四种条件：（C1）不使用生成式 AI 的 PjBL；（C2）使用无限制通用 LLM 的 PjBL；（C3）具有角色型辅导与协作、但没有 OpenPBL 权威编译器、证据门控和可撤销作品操作的 AI；（C4）完整 OpenPBL。C1 估计 AI 接入的贡献与成本，C2 检验即时生产力能否与独立学习分离，C3 将社会角色组织与可执行权威治理区分开来，C4 检验完整设计。随机化应在班级或团队层面进行以限制污染，并在分配和模型中处理基线成绩与教师效应。",
        "",
        "主要验证性结果是与课程目标一致的延迟、无 AI 迁移表现。次要结果包括即时后测、对条件盲评的作品质量、认识性能动性、认知卸载、证据链完整度和每组教师干预分钟数。关键对比是 C4 对 C3，用于检验可执行权威边界是否超越角色型 AI 产生额外价值；C4 对 C2 则检验护栏能否在不抹去生产性帮助的情况下保护迁移与证据。作品质量不得单独作为结论依据，因为流畅的 AI 生产可能掩盖较弱的独立理解 [3, 4]。",
        "",
        "### 7.2 分析与报告计划",
        "",
        "验证性模型应以条件预测延迟迁移，同时校正前测、班级聚类以及预先规定的教师和学科因素。次要模型采用同一组对比矩阵，并报告多重性控制、效应量与置信区间、缺失数据敏感性分析及意向治疗估计。过程分析可以按预注册方案检验“核验行为”和“证据链质量”的中介作用，但除非充分论证其假设，中介关系仍应解释为关联。公平性分析应报告先备知识与无障碍需求的异质性效应，但不得利用小样本亚组进行自动裁决。错误先修路由、未核验 AI 内容、伪造证据、不可访问交互、隐私事件和预警疲劳等不良结果，必须与收益一起报告。",
        "",
        "### 7.3 固定种子数据生成情景",
        "",
        f"为调试分析流程，我们使用 NumPy 的 PCG64 生成器和固定随机种子 {SEED} 生成 240 条合成记录，每个条件 60 条；各组前测来自同一分布。情景刻意编码相互重叠且并不完美的趋势，而非确定性优势：无限制 LLM 提高即时后测和作品质量，但增加认知卸载并削弱延迟迁移；角色型 AI 改善生产性结果，却只能部分保护能动性；完整 OpenPBL 改善延迟迁移和证据完整度，并减少教师干预时间。这些只是为了检验计划指标能否区分生产力、学习、能动性和课堂编排而设定的假设，并非观察效应。",
        "",
        *table3,
        "",
        *table4,
        "",
        "{{FIGURE_SYNTHETIC}}",
        "",
        "### 7.4 流程输出的解释",
        "",
        f"在合成流程试运行中，C1、C2、C3 和 C4 的延迟迁移均值（标准差）分别为 {fmt(summary, 'PBL-only', 'delayed_transfer')}、{fmt(summary, 'Unrestricted LLM', 'delayed_transfer')}、{fmt(summary, 'Role-based AI', 'delayed_transfer')} 和 {fmt(summary, 'OpenPBL full', 'delayed_transfer')}。无限制 LLM 条件的作品质量高于仅 PjBL 条件，却具有最低的能动性和证据完整度；完整 OpenPBL 条件具有最高的延迟迁移、能动性和证据完整度，以及最低的教师干预时间。该联合模式的价值在于使论文主张可以被证伪：真实数据完全可能显示控制机制只增加摩擦而没有学习收益，角色型 AI 已经足够，或 OpenPBL 的优势仅限于记录完整而不是学习。",
        "",
        "本文不报告 p 值，因为对虚构观察进行推断显著性检验没有意义。该试运行只有三项正当产出：冻结的数据结构和对比计划、确认分析能够区分预定构念，以及明确的替换清单。正式投稿前，研究者必须预注册方案、获得伦理审批、替换 CSV、重新生成表 3–4 和图 3、归档分析代码，并重写第 7.3–7.4 节及所有依赖占位数据的表述。",
        "",
    ])


def replace_citations(text: str, chinese: bool) -> str:
    mapping = ZH_CITATIONS if chinese else CITATIONS
    for source, target in sorted(mapping.items(), key=lambda item: len(item[0]), reverse=True):
        text = text.replace(source, target)
    return text


def build_markdown(source: Path, summary: pd.DataFrame, chinese: bool) -> str:
    text = source.read_text(encoding="utf-8")
    ref_heading = "## 参考文献" if chinese else "## References"
    body, _ = text.split(ref_heading, 1)
    if chinese:
        front = """# OpenPBL：将教学权威编译为证据门控的人机协同项目学习

**稿件类型：** Regular paper 稿件；系统设计与形成性评估

**格式说明：** 参考 JCST 常规论文的章节层级与数字引文；投稿场次确定后套用官方终稿模板

**作者信息：** 匿名审稿版本

**开发快照：** 2026 年 8 月 13 日

**研究状态说明：** 第 1–6 节报告系统设计与真实工程测试；第 7 节使用固定随机种子的合成占位数据，仅用于调试研究方案和分析流程，正式投稿前必须由经过伦理审批的真实数据完全替换。

## 摘要

生成式 AI 可以辅导，也可以直接参与作品生产，但这两类能力会遮蔽一个更基础的问题：谁控制课程、适应性支持、作品修改与高后果判断？本文提出 OpenPBL，一个把**教学权威视为类型化控制平面**的项目式学习环境。系统在教师、学生与 AI 之间分配课程、教学、适应性、生产和评价五类权威，并把分配结果编译为可执行机制：角色感知课程图区分真实先修与课内新授；“解释—示例—练习—反馈”之后只保留一次终结性掌握度测评；诊断仅面向被批准的先修点，按缺口一对一补救，所有学习者随后返回完整主课；AI 只能在学生明确委托下修改白名单草稿字段，并保留前后值、来源、冲突保护和撤销；项目推进依赖类型化的“测试—解释—修订”证据链，学生保留最终提交权，教师保留评价裁决权。我们通过架构决策记录、源码级理论—机制追踪和 14 个定向测试文件评估实现忠实度，111 项测试全部通过。为准备而非伪装课堂证据，本文进一步提出四组比较研究并在固定种子的合成数据上试运行其分析流程；这些数值被显著标注为非实证数据，不构成教育效果证据。本文贡献包括教学权威拓扑、从教育原则到软件权限的编译范式、端到端可审计实现，以及可证伪的实证检验方案。

**关键词：** 人工智能教育应用；项目式学习；教学权威；人机协作；学习者能动性；证据门控；可撤销委托；教育设计研究
"""
        body = re.sub(r"\A.*?\*\*关键词：\*\*.*?\n", front, body, count=1, flags=re.S)
        body = body.replace(
            "**RQ3：** 当前 OpenPBL 制品在多大程度上忠实实现了上述原则？在提出教育有效性主张之前，还必须补充哪些证据？",
            "**RQ3：** 当前 OpenPBL 制品在多大程度上忠实实现了上述原则？\n\n**RQ4：** 何种比较设计能够证伪“可执行权威边界比角色型或无限制 AI 更能保护独立迁移、能动性和教师判断”的主张？",
        )
        body = re.sub(
            r"本文(?:的贡献有四点|具有四项贡献)。.*?(?:能够检验上述机制的课堂研究方案|能够检验其作用机制的课堂研究方案)。",
            "本文的贡献有四点。第一，提出课程、教学、适应性、生产和评价五维教学权威拓扑。第二，把该拓扑形式化为具有行为者、状态前置条件、允许操作、不可变证据、撤销路径和最终人类裁决的控制平面。第三，展示如何把控制平面编译为角色化课程图、先修诊断循环、唯一终结性测评、可撤销工作区操作和证据门控状态转换。第四，严格区分实现忠实度与教育效果，并给出可证伪的四组比较方案。",
            body,
            count=1,
            flags=re.S,
        )
        related = """### 2.3 与智能体中心在线学习环境的关系

OpenMAIC [10] 等多智能体课程环境说明，协同的 LLM 角色可以连接课程创作与在线课堂活动。OpenPBL 选择了不同的设计单位：重点不是智能体角色的数量或覆盖范围，而是课程、适应性支持、作品生产、证据和评价中高后果权威的可执行分配。多智能体组织在本文中是一种实现选项，权威边界本身才是需要被设计和检验的研究对象。

"""
        body = body.replace("### 2.3 教育护栏必须是教学性的，而不仅是内容安全过滤", related + "### 2.4 教育护栏必须是教学性的，而不仅是内容安全过滤")
        body = body.replace("### 2.4 先备知识、适应性与课程效度", "### 2.5 先备知识、适应性与课程效度")
        body = body.replace("### 2.5 调节、证据与教师课堂编排", "### 2.6 调节、证据与教师课堂编排")
        body = body.replace("## 3. 概念框架：教育 AI 的权威拓扑", "## 3. 概念框架：作为可执行控制平面的教学权威")
        body = body.replace(
            "该框架从一个简单命题出发：**教育 AI 系统在很大程度上由‘谁可以依据什么做出何种高后果决定，以及该决定能否被争议或撤销’来界定**。本文区分五类权威（表 1）。",
            "权威拓扑只有在能够改变软件许可的行为时才具有设计价值。因此，我们把教学权威建模为类型化控制平面：它不直接生成课程内容或作品，而是规定哪些行为者在何种状态下可以执行哪些操作、必须留下何种证据、如何撤销或申诉，以及谁保留最终决定。一个权威规则只有同时声明行为者、状态前置条件、允许操作、不可变记录、逆转路径和终局人类裁决，才被视为可编译。表 1 给出五类权威及其边界。",
        )
        body = body.replace("## 5. 重构后的 OpenPBL 教育干预", "## 5. OpenPBL 系统架构与可执行机制")
        body = body.replace("## 6. 形成性忠实度结果", "## 6. 形成性制品评估")
        body = re.sub(
            r"(?<=### 6\.2 定向回归证据\n\n).*?(?=\n### 6\.3 忠实度结果能够说明什么、不能说明什么)",
            "本文在开发快照上执行了 14 个定向测试文件，覆盖课程入口生成、知识角色验证、适应性路由、深度交互排序、终结性测评、教学工具规划、工作区操作、伴学指导、学习证据准备度与阶段门控。测试命令于 2026 年 8 月 13 日运行，共 14/14 个测试文件、111/111 项测试通过，耗时 14.50 秒。\n\n这些结果证明当前快照符合被测试的教学契约，包括新授节点不得触发先修诊断、拓展不得插入终结性测评之前、AI 撤销不得覆盖已发生分歧的字段，以及单独上传作品不能满足制作阶段准备度。它们不代表全仓库验证、生产可靠性或教育效果。\n",
            body,
            count=1,
            flags=re.S,
        )
        body = body.replace("## 10. 结论", "## 11. 结论")
        body = body.replace("### 9.1 ", "### 10.1 ").replace("### 9.2 ", "### 10.2 ").replace("### 9.3 ", "### 10.3 ")
        body = body.replace("## 9. 实证研究议程", "## 10. 实证研究议程")
        body = body.replace("## 8. 局限与效度威胁", "## 9. 局限与效度威胁")
        body = body.replace("### 7.1 ", "### 8.1 ").replace("### 7.2 ", "### 8.2 ").replace("### 7.3 ", "### 8.3 ").replace("### 7.4 ", "### 8.4 ").replace("### 7.5 ", "### 8.5 ")
        body = body.replace("## 7. 讨论", "## 8. 讨论")
        body = body.replace("## 8. 讨论", synthetic_section(summary, True) + "\n## 8. 讨论")
        body = body.replace("第三，定向回归集合并不等于整个仓库测试套件，而且仍有一项预期失败。代码、模型版本、提示词和测试必须在实证部署前冻结并能够独立复现。", "第三，定向回归集合并不等于整个仓库测试套件；111 项通过只能说明被选择的边界在当前快照中符合规范。代码、模型版本、提示词、依赖环境与完整测试必须在实证部署前冻结并能够独立复现。")
        body = body.replace("专家评审还应明确检验零先修情形，以及当前 AI 主题默认结构之外的学科领域。", "专家评审还应明确检验“至少一个先修点”的硬性下限是否制造了虚假依赖，并覆盖当前 AI 主题默认结构之外的学科领域。")
        stage3 = """### 10.3 阶段三：比较效果研究

在共同设计和可行性标准达标后，可开展预注册的整群随机或严格匹配研究，比较四种条件：不使用生成式 AI 的 PjBL、使用无限制通用 LLM 的 PjBL、只有角色组织而没有权威编译与证据门控的 AI，以及完整 OpenPBL。主要结果应是与课程目标一致的延迟无 AI 迁移；次要结果包括概念成绩、作品质量、核验准确性、修订质量、AI 撤除后的坚持、认识性能动性、所有权、信任校准、协作质量、公平性与教师编排负担。C3 与 C4 的对比尤其关键，因为它把“角色丰富”与“可执行权威治理”分离开来。

不良结果必须被视作一等研究发现，包括复制但未核验的内容、独立表现下降、错误先修路由、证据伪造、不可访问的交互、有偏建议与教师预警疲劳。该议程检验的是一个可以失败的命题：有边界的权威控制也许比无限制帮助更能保护学习，但也可能只增加摩擦、文档负担或教师工作量。
"""
        body = re.sub(r"### 10\.3 阶段三：比较效果研究\n\n.*?(?=\n## 11\. 结论)", stage3, body, count=1, flags=re.S)
        body = re.sub(
            r"OpenPBL 试图回应 AIED 中的一个核心难题：.*?而不是已经被证明有效的教育解决方案。",
            "OpenPBL 回应的不是“再增加一种智能体角色”，而是生成式 AI 进入完整项目学习过程后，课程、教学、适应性、生产和评价权威如何保持可见、可执行并可追责。系统把教学原则编译为角色化课程图、严格先修补救、教学后唯一掌握度测评、白名单且可撤销的作品操作、类型化证据链，以及学生提交和教师裁决。\n\n形成性评估表明，这些边界可以成为可检查的软件契约；14 个定向测试文件中的 111 项测试全部通过。但工程符合性不是学习效果。第 7 节的合成数值只证明拟定的数据结构和分析流程能够运行，并不证明任何条件更优。OpenPBL 当前最有根据的定位，是一种围绕教学权威控制平面的原创、可证伪系统设计；只有在教师共同设计、课堂可行性研究和预注册比较试验完成后，才能提出关于学习、能动性或工作量的效果结论。",
            body,
            count=1,
            flags=re.S,
        )
        availability = """## 数据与制品可获得性

本文分析对象为 2026 年 8 月 13 日的 OpenPBL 开发快照。14 个定向测试文件中的 111 项测试全部通过。第 7 节的 240 条记录使用固定种子 20260813 生成，并在每行标记 `synthetic_placeholder=true`；它们仅用于流程调试。真实课堂研究开始前应冻结代码修订、数据库结构、架构决策记录、提示词与模型配置、测试命令、分析脚本和数据字典。本文未收集人类参与者数据。

## 作者贡献、资助与利益冲突

匿名审稿阶段暂不披露。非匿名版本应按 CRediT 分类报告贡献，列出全部资助，并披露相关财务或非财务利益。
"""
        body = body.split("## 数据与制品可获得性", 1)[0].rstrip() + "\n\n" + availability
    else:
        front = """# OpenPBL: Compiling Pedagogical Authority into Evidence-Gated Human-AI Project Learning

**Manuscript type:** Regular-paper draft; system design and formative evaluation

**Format note:** JCST-inspired section hierarchy and numeric citations; apply the official publisher template after venue confirmation

**Author information:** Anonymous for peer review

**Development snapshot:** 13 August 2026

**Research-status notice:** Sections 1-6 report system design and real engineering tests. Section 7 uses seeded synthetic placeholders solely to debug the study protocol and analysis pipeline; approved empirical data must replace them before submission.

## Abstract

Generative AI can tutor and co-produce, but those capabilities obscure a more basic question: who controls curriculum, adaptation, artifact changes, and consequential judgments? We present OpenPBL, a project-based learning environment that treats **pedagogical authority as a typed control plane**. Five authority domains—curricular, instructional, adaptive, productive, and evaluative—are allocated across teachers, students, and AI and compiled into executable mechanisms. A role-aware curriculum graph distinguishes authentic prerequisites from lesson targets; instruction follows explanation, example, practice, and feedback before one terminal mastery assessment; diagnostics target only approved prerequisites, route one-to-one remediation, and return every learner to the full lesson; AI may edit whitelisted draft fields only under explicit student delegation with before/after values, provenance, conflict protection, and undo; and progression requires typed test-interpret-revise evidence while students retain submission and teachers retain adjudication. We evaluate implementation fidelity through architecture-decision records, source-level theory-to-mechanism traceability, and 14 targeted test files; all 111 tests passed. To prepare rather than imitate classroom evidence, we also specify a four-condition comparative study and dry-run its pipeline on a seeded synthetic dataset, prominently labelled non-empirical. The synthetic output is not evidence of learning effects. Contributions are an authority topology, a compilation pattern from pedagogical principles to software permissions, an end-to-end auditable implementation, and a falsifiable empirical protocol.

**Keywords:** artificial intelligence in education; project-based learning; pedagogical authority; human-AI collaboration; learner agency; evidence gates; reversible delegation; educational design research
"""
        body = re.sub(r"\A.*?\*\*Keywords:\*\*.*?\n", front, body, count=1, flags=re.S)
        body = body.replace(
            "**RQ3.** To what extent does the implemented OpenPBL artifact faithfully instantiate these principles, and what remains to be established before claims about educational effectiveness are warranted?",
            "**RQ3.** To what extent does the implemented OpenPBL artifact faithfully instantiate these principles?\n\n**RQ4.** What comparative design can falsify the claim that executable authority boundaries better preserve independent transfer, agency, and teacher judgment than role-based or unrestricted AI?",
        )
        body = re.sub(
            r"The paper makes four contributions\..*?capable of testing the proposed mechanisms\.",
            "The paper makes four contributions. First, it proposes a five-domain topology of curricular, instructional, adaptive, productive, and evaluative authority. Second, it formalizes that topology as a control plane with actors, state preconditions, permitted operations, immutable evidence, reversal paths, and terminal human decisions. Third, it shows how the control plane compiles into a role-aware curriculum graph, prerequisite loop, single terminal assessment, reversible workspace operations, and evidence-gated state transitions. Fourth, it separates implementation fidelity from educational effect and specifies a falsifiable four-condition comparison.",
            body,
            count=1,
            flags=re.S,
        )
        related = """### 2.3 Relation to agent-centric online learning environments

Multi-agent course environments such as OpenMAIC [10] show that coordinated LLM roles can connect course authoring with online classroom activity. OpenPBL takes a different unit of design: not the number or coverage of agent roles, but the executable allocation of consequential authority across curriculum, adaptation, artifact production, evidence, and evaluation. Multi-agent organization is therefore an implementation option in this paper; authority boundaries are the object to be designed and tested.

"""
        body = body.replace("### 2.3 Guardrails must be pedagogical, not merely safety filters", related + "### 2.4 Guardrails must be pedagogical, not merely safety filters")
        body = body.replace("### 2.4 Prior knowledge, adaptation, and curricular validity", "### 2.5 Prior knowledge, adaptation, and curricular validity")
        body = body.replace("### 2.5 Regulation, evidence, and teacher orchestration", "### 2.6 Regulation, evidence, and teacher orchestration")
        body = body.replace("## 3. Conceptual Framework: An Authority Topology for Educational AI", "## 3. Conceptual Framework: Pedagogical Authority as an Executable Control Plane")
        body = body.replace(
            "The framework begins from a simple proposition: **an educational AI system is partly defined by who may make which consequential decision, on what basis, and whether that decision can be contested or reversed**. We distinguish five forms of authority (Table 1).",
            "An authority topology becomes useful only when it changes what software permits. We therefore model pedagogical authority as a typed control plane: it does not itself generate lessons or artifacts; it specifies which actor may perform which operation under which state, what evidence must be preserved, how a decision can be reversed or contested, and who retains the terminal judgment. An authority rule is compileable only when it declares an actor, state precondition, permitted operation, immutable record, reversal path, and final human decision. Table 1 applies this grammar to five domains.",
        )
        body = body.replace("## 5. The Redesigned OpenPBL Intervention", "## 5. OpenPBL System Architecture and Executable Mechanisms")
        body = body.replace("## 6. Formative Fidelity Results", "## 6. Formative Artifact Evaluation")
        body = re.sub(
            r"Fourteen targeted test files covering course-entry generation.*?rather than hidden in interface behaviour\.",
            "Fourteen targeted test files covering course-entry generation, knowledge-role validation, adaptive routing, deep-interaction sequencing, terminal assessment, teaching-tool planning, workspace operations, companion guidance, learning-evidence readiness, and stage gates were executed against the development snapshot. The run on 13 August 2026 completed with 14/14 test files and 111/111 tests passing in 14.50 seconds.\n\nThe result demonstrates conformance of the tested snapshot to selected pedagogical contracts: lesson nodes cannot trigger prerequisite diagnosis; enrichment cannot precede the terminal assessment; undo cannot overwrite a diverged field; and artifact upload alone cannot satisfy making-stage readiness. It is neither a full-repository verification nor evidence of production reliability or educational effect.",
            body,
            count=1,
            flags=re.S,
        )
        body = body.replace("## 10. Conclusion", "## 11. Conclusion")
        body = body.replace("### 9.1 ", "### 10.1 ").replace("### 9.2 ", "### 10.2 ").replace("### 9.3 ", "### 10.3 ")
        body = body.replace("## 9. Empirical Research Agenda", "## 10. Empirical Research Agenda")
        body = body.replace("## 8. Limitations and Threats to Validity", "## 9. Limitations and Threats to Validity")
        body = body.replace("### 7.1 ", "### 8.1 ").replace("### 7.2 ", "### 8.2 ").replace("### 7.3 ", "### 8.3 ").replace("### 7.4 ", "### 8.4 ").replace("### 7.5 ", "### 8.5 ")
        body = body.replace("## 7. Discussion", "## 8. Discussion")
        body = body.replace("## 8. Discussion", synthetic_section(summary, False) + "\n## 8. Discussion")
        body = body.replace("Third, the targeted regression set is not the entire repository suite, and one expectation failed. The code, model versions, prompts, and tests must be frozen and independently reproducible before empirical deployment.", "Third, the targeted regression set is not the entire repository suite; 111 passing tests show only that selected boundaries conform in the current snapshot. The code revision, dependency environment, model versions, prompts, and complete tests must be frozen and independently reproducible before empirical deployment.")
        body = body.replace("Expert review should explicitly test zero-prerequisite cases and domains outside the current AI-oriented defaults.", "Expert review should explicitly test whether the hard minimum of one prerequisite manufactures dependencies, as well as domains outside the current AI-oriented defaults.")
        stage3 = """### 10.3 Stage 3: Comparative effectiveness study

After co-design and feasibility criteria are met, a preregistered cluster-randomized or carefully matched study can compare four conditions: PjBL without generative AI, PjBL with an unrestricted general-purpose LLM, role-based AI without authority compilation or evidence gates, and full OpenPBL. The primary outcome should be delayed, unaided transfer aligned with lesson targets. Secondary outcomes should include conceptual achievement, artifact quality, verification accuracy, revision quality, persistence after AI withdrawal, epistemic agency, ownership, trust calibration, collaboration quality, equity, and teacher orchestration workload. The C3-versus-C4 contrast is essential because it separates role richness from executable authority governance.

Adverse outcomes should be treated as first-class results: copied but unverified content, reduced unaided performance, false prerequisite routing, evidence fabrication, inaccessible interactions, biased suggestions, privacy incidents, and teacher alert fatigue. The agenda tests a proposition that can fail: bounded authority control may protect learning better than unrestricted assistance, but it may instead add friction, documentation burden, or teacher work.
"""
        body = re.sub(r"### 10\.3 Stage 3: Comparative effectiveness study\n\n.*?(?=\n## 11\. Conclusion)", stage3, body, count=1, flags=re.S)
        body = re.sub(
            r"OpenPBL addresses a central AIED problem:.*?not a proven educational solution\.",
            "OpenPBL addresses more than the addition of another agent role. Once generative AI enters an end-to-end project-learning process, curricular, instructional, adaptive, productive, and evaluative authority must remain visible, executable, and accountable. The system compiles pedagogical principles into a role-aware curriculum graph, strict prerequisite remediation, one post-instruction mastery assessment, whitelisted and reversible artifact operations, typed evidence chains, student submission, and teacher adjudication.\n\nThe formative evaluation shows that these boundaries can become inspectable software contracts: all 111 tests across 14 targeted files passed. Engineering conformance is not learning effect. The synthetic values in Section 7 demonstrate only that the proposed data schema and analysis pipeline run; they establish no condition as superior. OpenPBL's warranted status is therefore an original and falsifiable system design organized around a pedagogical-authority control plane. Claims about learning, agency, or workload require teacher co-design, classroom feasibility evidence, and a preregistered comparative trial.",
            body,
            count=1,
            flags=re.S,
        )
        availability = """## Data and Artifact Availability

The analyzed artifact is the OpenPBL development snapshot of 13 August 2026. All 111 tests in 14 targeted files passed. The 240 records in Section 7 were generated with fixed seed 20260813 and carry `synthetic_placeholder=true` on every row; they exist only to debug the pipeline. Before classroom research, the authors should freeze the source revision, schema, architecture-decision records, prompt and model configuration, dependency environment, test commands, analysis code, and data dictionary. No human-participant data were collected.

## Author Contributions, Funding, and Competing Interests

These statements are withheld for anonymous review. The non-anonymous version should report contributions using the CRediT taxonomy, identify all funding, and disclose relevant financial or non-financial interests.
"""
        body = body.split("## Data and Artifact Availability", 1)[0].rstrip() + "\n\n" + availability
    body = replace_citations(body, chinese)
    return body.rstrip() + "\n\n" + ref_heading + "\n\n" + "\n\n".join(REFERENCES) + "\n"


def parse_markdown(text: str) -> list[Token]:
    lines = text.splitlines()
    tokens: list[Token] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            tokens.append(Token("paragraph", " ".join(item.strip() for item in paragraph)))
            paragraph.clear()

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            flush()
            i += 1
            continue
        if stripped.startswith("{{") and stripped.endswith("}}"):
            flush()
            tokens.append(Token("figure", stripped))
            i += 1
            continue
        if stripped.startswith("|"):
            flush()
            raw: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                raw.append(lines[i].strip())
                i += 1
            rows: list[list[str]] = []
            for idx, line in enumerate(raw):
                values = [item.strip() for item in line.strip("|").split("|")]
                if idx == 1 and all(re.fullmatch(r":?-{3,}:?", item) for item in values):
                    continue
                rows.append(values)
            tokens.append(Token("table", rows))
            continue
        if stripped.startswith("### "):
            flush()
            tokens.append(Token("h2", stripped[4:]))
            i += 1
            continue
        if stripped.startswith("## "):
            flush()
            tokens.append(Token("h1", stripped[3:]))
            i += 1
            continue
        if stripped.startswith("# "):
            flush()
            tokens.append(Token("title", stripped[2:]))
            i += 1
            continue
        if stripped.startswith("- "):
            flush()
            tokens.append(Token("bullet", stripped[2:]))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            flush()
            tokens.append(Token("number", stripped))
            i += 1
            continue
        paragraph.append(stripped)
        i += 1
    flush()
    return tokens


def set_cell_width(cell, inches: float) -> None:
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def table_widths(rows: list[list[str]], total: float) -> list[float]:
    cols = len(rows[0])
    lengths = []
    for col in range(cols):
        maximum = max(len(re.sub(r"[*`]", "", row[col])) for row in rows)
        lengths.append(max(8, min(maximum, 42)))
    if cols >= 5:
        lengths[0] = max(lengths[0], 18)
    raw = [max(0.72, total * length / sum(lengths)) for length in lengths]
    scale = total / sum(raw)
    return [value * scale for value in raw]


def add_table(doc: Document, rows: list[list[str]], font_name: str, size: float, total: float, bilingual: bool = False, zh_rows: list[list[str]] | None = None) -> None:
    widths = table_widths(rows, total)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        legacy._keep_row_together(table.rows[r_idx])
        if r_idx == 0:
            legacy._repeat_header(table.rows[r_idx])
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_width(cell, widths[c_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            legacy._cell_margins(cell, top=70, start=80, bottom=70, end=80)
            if r_idx == 0:
                legacy._shade(cell, INK)
            elif r_idx % 2 == 0:
                legacy._shade(cell, PALE_GRAY)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.keep_with_next = r_idx == 0
            legacy._inline(p, value, font_name, max(7.2, size - 1.3))
            if bilingual and zh_rows is not None:
                zhp = cell.add_paragraph()
                zhp.paragraph_format.space_before = Pt(2)
                zhp.paragraph_format.space_after = Pt(0)
                zhp.paragraph_format.line_spacing = 1.0
                zhp.paragraph_format.keep_with_next = r_idx == 0
                legacy._inline(zhp, zh_rows[r_idx][c_idx], "SimSun", max(7.1, size - 1.5))
                for run in zhp.runs:
                    run.font.color.rgb = RGBColor.from_string("40566A")
            if r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def configure(doc: Document, bilingual: bool) -> tuple[str, float, float]:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    margin = 0.68 if bilingual else 0.74
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.32)
    font_name = "Times New Roman"
    size = 9.35 if bilingual else 10.15
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun" if bilingual else font_name)
    normal.font.size = Pt(size)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.13 if bilingual else 1.16
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.widow_control = True
    for level, hsize in ((1, 14.3), (2, 11.6)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun" if bilingual else font_name)
        style.font.size = Pt(hsize)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK if level == 1 else BLUE)
        style.paragraph_format.space_before = Pt(9 if level == 1 else 6)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    running = "OPENPBL | JCST-STYLE REGULAR-PAPER DRAFT"
    if bilingual:
        running += " | 中英对照审阅稿"
    legacy._set_run_font(header.add_run(running), font_name, 7.2, bold=True, color=MUTED)
    legacy._page_number(section.footer.paragraphs[0])
    for run in section.footer.paragraphs[0].runs:
        legacy._set_run_font(run, font_name, 7.5, color=MUTED)
    return font_name, size, 8.27 - 2 * margin


def add_callout(doc: Document, en: str, zh: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    legacy._shade(cell, PALE_RED)
    legacy._cell_margins(cell, top=115, start=130, bottom=115, end=130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    legacy._inline(p, en, "Times New Roman", 9.0)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(SYNTH_RED)
    if zh:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(3)
        p2.paragraph_format.space_after = Pt(0)
        legacy._inline(p2, zh, "SimSun", 8.8)
        for run in p2.runs:
            run.font.color.rgb = RGBColor.from_string(SYNTH_RED)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_title(doc: Document, en: str, zh: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    legacy._set_run_font(p.add_run(en), "Times New Roman", 18.2 if not zh else 17.0, bold=True, color=INK)
    if zh:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(7)
        p2.paragraph_format.keep_with_next = True
        legacy._set_run_font(p2.add_run(zh), "SimSun", 15.5, bold=True, color=TEAL)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(5)
    ppr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), BLUE)
    pbdr.append(bottom)
    ppr.append(pbdr)


def add_figure(doc: Document, marker: str, bilingual: bool) -> None:
    mapping = {
        "{{FIGURE_COAGENCY}}": (
            ASSETS_V2 / "authority_topology_en.png",
            ASSETS_V2 / "authority_topology_zh.png",
            "Figure 1. Five-domain authority topology and executable safeguards.",
            "图 1. 五维教学权威拓扑及其可执行护栏。",
        ),
        "{{FIGURE_SEQUENCE}}": (
            ASSETS_V2 / "learning_sequence_en.png",
            ASSETS_V2 / "learning_sequence_zh.png",
            "Figure 2. Sequence from curricular boundaries and adaptive teaching to evidence-gated project collaboration.",
            "图 2. 从课程边界和适应性教学到证据门控项目协作的完整序列。",
        ),
        "{{FIGURE_SYNTHETIC}}": (
            ASSETS_V3 / "synthetic_outcomes_en.png",
            ASSETS_V3 / "synthetic_outcomes_zh.png",
            "Figure 3. SYNTHETIC PLACEHOLDER scenario used only to dry-run the analysis pipeline.",
            "图 3. 仅用于试运行分析流程的合成占位情景。",
        ),
    }
    en_path, zh_path, en_cap, zh_cap = mapping[marker]
    paths = [en_path]
    if bilingual:
        paths = [en_path, zh_path] if marker != "{{FIGURE_SYNTHETIC}}" else [zh_path]
    for path in paths:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        width = 6.74 if not bilingual else 6.62
        p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    legacy._set_run_font(cap.add_run(en_cap), "Times New Roman", 8.2, italic=True, color=SYNTH_RED if "SYNTHETIC" in en_cap else MUTED)
    if bilingual:
        legacy._set_run_font(cap.add_run("\n" + zh_cap), "SimSun", 8.0, italic=True, color=SYNTH_RED if "合成" in zh_cap else MUTED)


def is_notice(text: str) -> bool:
    lowered = text.lower()
    return "research-status notice" in lowered or "synthetic placeholder—not empirical" in lowered or "研究状态说明" in text or "合成占位数据——不是实证结果" in text


def add_paragraph(doc: Document, text: str, font: str, size: float, reference: bool = False, color: str = BLACK) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.14
    p.paragraph_format.space_after = Pt(4)
    if reference:
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.keep_together = True
        size = min(size, 8.65)
    legacy._inline(p, text, font, size)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)


def build_english(markdown: str, output: Path) -> None:
    tokens = parse_markdown(markdown)
    doc = Document()
    font, size, text_width = configure(doc, bilingual=False)
    in_references = False
    metadata = True
    for token in tokens:
        if token.kind == "title":
            add_title(doc, str(token.value))
        elif token.kind == "h1":
            heading = str(token.value)
            metadata = False
            if heading == "References":
                doc.add_page_break()
                in_references = True
            elif in_references:
                in_references = False
            p = doc.add_paragraph(style="Heading 1")
            legacy._inline(p, heading, font, 14.3)
        elif token.kind == "h2":
            caption = str(token.value)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            if caption.startswith("Table"):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                legacy._set_run_font(p.add_run(caption), font, 9.0, bold=True, color=SYNTH_RED if "SYNTHETIC" in caption else INK)
            else:
                p.style = doc.styles["Heading 2"]
                legacy._inline(p, caption, font, 11.6)
        elif token.kind == "paragraph":
            value = str(token.value)
            if is_notice(value):
                add_callout(doc, value)
            elif metadata:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                legacy._inline(p, value, font, 8.6)
            else:
                add_paragraph(doc, value, font, size, reference=in_references)
        elif token.kind == "table":
            add_table(doc, token.value, font, size, text_width - 0.05)
        elif token.kind == "figure":
            add_figure(doc, str(token.value), bilingual=False)
        elif token.kind in ("bullet", "number"):
            p = doc.add_paragraph(style="List Bullet" if token.kind == "bullet" else None)
            p.paragraph_format.left_indent = Inches(0.30)
            p.paragraph_format.first_line_indent = Inches(-0.20)
            p.paragraph_format.space_after = Pt(2.5)
            legacy._inline(p, str(token.value), font, size)
    props = doc.core_properties
    props.title = str(tokens[0].value)
    props.subject = "JCST-inspired regular-paper draft: system design, formative evaluation, and synthetic protocol dry run"
    props.author = "Anonymous for peer review"
    props.keywords = "AIED; project-based learning; pedagogical authority; human-AI collaboration"
    doc.save(output)


def build_bilingual(en_markdown: str, zh_markdown: str, output: Path) -> None:
    en_tokens = parse_markdown(en_markdown)
    zh_tokens = parse_markdown(zh_markdown)
    if len(en_tokens) != len(zh_tokens):
        raise ValueError(f"Token-count mismatch: EN={len(en_tokens)} ZH={len(zh_tokens)}")
    for idx, (en, zh) in enumerate(zip(en_tokens, zh_tokens)):
        if en.kind != zh.kind:
            raise ValueError(f"Token-kind mismatch at {idx}: {en.kind} != {zh.kind}")
        if en.kind == "table":
            if len(en.value) != len(zh.value) or any(len(a) != len(b) for a, b in zip(en.value, zh.value)):
                raise ValueError(f"Table-shape mismatch at token {idx}")
    doc = Document()
    font, size, text_width = configure(doc, bilingual=True)
    in_references = False
    metadata = True
    for en, zh in zip(en_tokens, zh_tokens):
        if en.kind == "title":
            add_title(doc, str(en.value), str(zh.value))
        elif en.kind == "h1":
            en_heading, zh_heading = str(en.value), str(zh.value)
            metadata = False
            if en_heading == "References":
                in_references = True
            elif in_references:
                in_references = False
            p = doc.add_paragraph(style="Heading 1")
            legacy._inline(p, en_heading, font, 13.5)
            legacy._set_run_font(p.add_run("  /  " + zh_heading), "SimSun", 12.4, bold=True, color=TEAL)
        elif en.kind == "h2":
            en_text, zh_text = str(en.value), str(zh.value)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            if en_text.startswith("Table"):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                color = SYNTH_RED if "SYNTHETIC" in en_text else INK
                legacy._set_run_font(p.add_run(en_text), font, 8.6, bold=True, color=color)
                legacy._set_run_font(p.add_run("\n" + zh_text), "SimSun", 8.4, bold=True, color=color)
            else:
                p.style = doc.styles["Heading 2"]
                legacy._inline(p, en_text, font, 10.9)
                legacy._set_run_font(p.add_run("  /  " + zh_text), "SimSun", 10.3, bold=True, color=TEAL)
        elif en.kind == "paragraph":
            en_text, zh_text = str(en.value), str(zh.value)
            if is_notice(en_text) or is_notice(zh_text):
                add_callout(doc, en_text, zh_text)
            elif metadata:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(1)
                legacy._inline(p, en_text, font, 8.0)
                legacy._set_run_font(p.add_run("\n" + zh_text), "SimSun", 7.9, color="40566A")
            elif in_references:
                add_paragraph(doc, en_text, font, 8.3, reference=True)
            else:
                add_paragraph(doc, en_text, font, size)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.08
                p.paragraph_format.space_after = Pt(5)
                p.paragraph_format.left_indent = Inches(0.10)
                legacy._inline(p, zh_text, "SimSun", size - 0.15)
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string("40566A")
        elif en.kind == "table":
            add_table(doc, en.value, font, size, text_width - 0.05, bilingual=True, zh_rows=zh.value)
        elif en.kind == "figure":
            add_figure(doc, str(en.value), bilingual=True)
        elif en.kind in ("bullet", "number"):
            p = doc.add_paragraph(style="List Bullet" if en.kind == "bullet" else None)
            p.paragraph_format.left_indent = Inches(0.30)
            p.paragraph_format.first_line_indent = Inches(-0.20)
            p.paragraph_format.space_after = Pt(1)
            legacy._inline(p, str(en.value), font, size)
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.38)
            p2.paragraph_format.space_after = Pt(3)
            legacy._inline(p2, str(zh.value), "SimSun", size - 0.2)
            for run in p2.runs:
                run.font.color.rgb = RGBColor.from_string("40566A")
    props = doc.core_properties
    props.title = "OpenPBL bilingual paragraph-aligned review manuscript"
    props.subject = "English-Chinese review copy; synthetic outcome section visibly labelled"
    props.author = "Anonymous for peer review"
    doc.save(output)


def audit_markdown(en: str, zh: str) -> dict[str, object]:
    en_body = en.split("## References", 1)[0]
    checks = {
        "english_word_count": len(re.findall(r"\b[\w'-]+\b", en_body)),
        "reference_count": len(REFERENCES),
        "openmaic_body_mentions": en_body.count("OpenMAIC"),
        "synthetic_warning_mentions_en": en_body.lower().count("synthetic placeholder"),
        "synthetic_warning_mentions_zh": zh.count("合成占位"),
        "old_test_counts_absent": not re.search(r"\b(?:113|114)\b", en_body),
        "old_test_counts_absent_zh": not re.search(r"\b(?:113|114)\b", zh.split("## 参考文献", 1)[0]),
        "targeted_test_claim_present": "111/111" in en_body and "14/14" in en_body,
        "author_year_citations_absent": not re.search(r"\([A-Z][^()]{0,80},\s*(?:19|20)\d{2}\)", en_body),
        "all_reference_numbers_present": all(
            re.search(rf"(?<!\d){re.escape(f'[{idx}]')}(?!\d)", en_body)
            or re.search(rf"\[(?:\d+,\s*)*{idx}(?:,\s*\d+)*\]", en_body)
            for idx in range(1, len(REFERENCES) + 1)
        ),
        "no_unlabelled_results_heading": "## 7. Results" not in en,
    }
    failures = [key for key, value in checks.items() if isinstance(value, bool) and not value]
    if checks["openmaic_body_mentions"] != 1:
        failures.append("openmaic_body_mentions_exactly_one")
    if checks["synthetic_warning_mentions_en"] < 3 or checks["synthetic_warning_mentions_zh"] < 3:
        failures.append("synthetic_warning_repetition")
    if checks["english_word_count"] < 6500:
        failures.append("english_word_count_at_least_6500")
    checks["failures"] = failures
    if failures:
        raise ValueError("Markdown audit failed: " + ", ".join(failures))
    return checks


def write_research_evidence(audit: dict[str, object]) -> None:
    analysis = """# Baseline writing analysis and originality transfer

## Baseline narrative craft

The reference article uses a clear situation-complication-answer progression: scalable online learning creates an opportunity; fragmented personalization and role coordination create a complication; a unified agent environment is proposed and then decomposed into system, technical, and empirical evaluations. Its strongest transferable techniques are an early conceptual figure, explicit subsystem boundaries, a separation of technical evaluation from educational evaluation, quantified methods, and bounded future-work claims.

## Limits not carried into OpenPBL

The baseline's breadth can dilute the mechanism under study, and observational or correlational evaluation cannot isolate which agent role causes an outcome. OpenPBL therefore does not compete on the number of agents, subjects, or platform features. It treats executable allocation of consequential authority as the independent conceptual contribution and mentions the baseline once in related work.

## OpenPBL argument

Situation: generative AI can tutor and co-produce throughout PjBL. Complication: role labels do not determine who controls curriculum, adaptation, artifact change, evidence, or evaluation. Answer: compile five pedagogical authority domains into typed permissions, state preconditions, immutable evidence, reversible operations, and terminal human decisions. Evaluation is split into real implementation fidelity and a visibly non-empirical protocol dry run.
"""
    (EVIDENCE_DIR / "baseline_writing_analysis.md").write_text(analysis, encoding="utf-8")
    manifest = {
        "development_snapshot": "2026-08-13",
        "paper_style": "JCST-inspired regular-paper review manuscript; not an official publisher template",
        "real_engineering_evidence": {
            "test_files": 14,
            "tests_passed": 111,
            "tests_failed": 0,
            "duration_seconds": 14.50,
            "scope": "targeted suite only",
        },
        "synthetic_placeholder": {
            "seed": SEED,
            "records": 240,
            "conditions": 4,
            "empirical_claim_allowed": False,
        },
        "audit": audit,
    }
    (EVIDENCE_DIR / "run_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    _, summary = generate_synthetic_data()
    draw_synthetic_figure(summary, ASSETS_V3 / "synthetic_outcomes_en.png", chinese=False)
    draw_synthetic_figure(summary, ASSETS_V3 / "synthetic_outcomes_zh.png", chinese=True)
    en = build_markdown(EN_SOURCE, summary, chinese=False)
    zh = build_markdown(ZH_SOURCE, summary, chinese=True)
    audit = audit_markdown(en, zh)
    EN_V3.write_text(en, encoding="utf-8")
    ZH_V3.write_text(zh, encoding="utf-8")
    write_research_evidence(audit)
    build_english(en, EN_DOCX)
    build_bilingual(en, zh, BI_DOCX)
    print(json.dumps({
        "english_markdown": str(EN_V3),
        "chinese_markdown": str(ZH_V3),
        "english_docx": str(EN_DOCX),
        "bilingual_docx": str(BI_DOCX),
        "audit": audit,
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
